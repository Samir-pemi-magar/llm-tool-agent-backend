"""
    The actual python-docx logic. This file lives ONLY inside the sandbox image —
    it never runs in the main app process.
"""

import docx


def read_docx_text(file_path: str) -> dict:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return {"num_paragraphs": len(paragraphs), "paragraphs": paragraphs}


def get_docx_tables(file_path: str) -> dict:
    doc = docx.Document(file_path)
    tables = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)
    return {"num_tables": len(tables), "tables": tables}


def find_and_replace_text(file_path: str, find: str, replace: str) -> dict:
    doc = docx.Document(file_path)
    replacements = 0
    for paragraph in doc.paragraphs:
        if find in paragraph.text:
            for run in paragraph.runs:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    replacements += 1
    doc.save(file_path)
    return {"status": "ok", "replacements": replacements}


REGISTRY = {
    "read_docx_text": read_docx_text,
    "get_docx_tables": get_docx_tables,
    "find_and_replace_text": find_and_replace_text,
}